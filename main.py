"""
A006 — Generador de Informes de Análisis Hidráulico
Proyecto URBITIK — Consorcio de Aguas de Álava (Urbide)

Uso: python main.py
"""
import os
import re
import sys
import logging
from pathlib import Path

# Mostrar info/warnings de los módulos word/ y ai/ por consola.
logging.basicConfig(
    level=logging.INFO,
    format="%(levelname)s [%(name)s] %(message)s",
)

from db.connection import test_connection
from db.queries import (
    get_municipio_nombre,
    get_result_ids_disponibles,
    get_datos_municipio,
    get_datos_sectores,
    get_resultados_simulacion,
)


# Campos que el usuario introduce manualmente porque no están en la BD
CAMPOS_MANUALES = [
    ("habitantes", "Población abastecida (habitantes)"),
    ("topografia", "Topografía del municipio (llana/ondulada/montañosa)"),
    ("fuentes_abastecimiento", "Fuentes de abastecimiento (texto breve)"),
    ("ubicacion_depositos", "Ubicación de los depósitos (texto breve)"),
    ("descripcion_bombeos", "Descripción de los bombeos (texto breve)"),
    ("abonados_domesticos", "Número de abonados domésticos"),
    ("factores_estacionales", "Factores estacionales (texto breve)"),
    ("periodo_demanda", "Periodo de referencia de demanda (ej: 16/03/2025-23/03/2025)"),
    ("etap", "Estación de tratamiento (ETAP) — nombre y descripción breve"),
]


PLANTILLAS_DIR = Path(__file__).resolve().parent / "plantillas"


def _validar_docx(ruta: Path) -> tuple[bool, str]:
    """Comprueba que el archivo es un .docx válido y abrible.

    Devuelve (ok, motivo). Detecta casos típicos de OneDrive On-Demand,
    archivos bloqueados por Word o ficheros corruptos.
    """
    if not ruta.exists():
        return False, f"No se encontró el archivo: {ruta}"
    if ruta.is_dir():
        return False, "La ruta es un directorio, no un archivo .docx."
    if ruta.suffix.lower() != ".docx":
        return False, f"El archivo no tiene extensión .docx ({ruta.suffix})."
    try:
        size = ruta.stat().st_size
    except OSError as e:
        return False, f"No se puede leer la metadata del archivo: {e}"
    if size == 0:
        return False, "El archivo tiene 0 bytes (placeholder de OneDrive sin descargar?)."
    # Un .docx es un ZIP. Comprobamos la firma "PK" en los primeros 2 bytes.
    try:
        with open(ruta, "rb") as f:
            head = f.read(4)
    except PermissionError:
        return False, "Permiso denegado. ¿Lo tienes abierto en Word? Ciérralo y reintenta."
    except OSError as e:
        return False, f"No se puede abrir el archivo: {e}"
    if not head.startswith(b"PK"):
        return False, (
            "El archivo no es un ZIP/.docx válido (¿descarga incompleta de OneDrive, "
            "o archivo corrupto?)."
        )
    return True, "ok"


def _listar_plantillas() -> list[Path]:
    """Devuelve la lista ordenada de .docx en la carpeta `plantillas/`."""
    if not PLANTILLAS_DIR.exists():
        return []
    return sorted(PLANTILLAS_DIR.glob("*.docx"))


def pedir_ruta_plantilla():
    """Solicita la plantilla Word.

    1) Si hay .docx en la carpeta `plantillas/`, los lista numerados y permite
       elegir uno (recomendado: ruta estable, sin OneDrive de por medio).
    2) Permite también introducir una ruta manual con la opción 'm'.
    Valida que sea un .docx abrible (no un placeholder de OneDrive ni un
    archivo bloqueado).
    """
    while True:
        plantillas_locales = _listar_plantillas()

        if plantillas_locales:
            print(f"\n  Plantillas disponibles en {PLANTILLAS_DIR.name}/:")
            for i, p in enumerate(plantillas_locales, 1):
                print(f"    [{i}] {p.name}")
            print("    [m] Introducir ruta manual")
            opcion = input("  Selecciona una opción: ").strip().lower()

            if opcion == "m":
                pass  # cae al flujo manual abajo
            elif opcion.isdigit() and 1 <= int(opcion) <= len(plantillas_locales):
                ruta = plantillas_locales[int(opcion) - 1]
                ok, motivo = _validar_docx(ruta)
                if ok:
                    return str(ruta)
                print(f"  ERROR: {motivo}")
                continue
            else:
                print("  ERROR: opción no válida.")
                continue
        else:
            print(f"\n  (No hay .docx en {PLANTILLAS_DIR}/. Introduce ruta manual.)")
            print("  Sugerencia: copia la plantilla a esa carpeta para no depender")
            print("  de rutas externas (OneDrive, unidades de red, etc.).")

        # Flujo manual
        raw = input("  Ruta completa de la plantilla (.docx): ").strip()
        if not raw:
            print("  ERROR: ruta vacía.")
            continue

        # Limpia comillas dobles, simples y curly quotes habituales al copiar
        for ch in ('"', "'", "“", "”", "‘", "’"):
            raw = raw.strip(ch)
        raw = raw.strip()

        ruta = Path(os.path.expandvars(os.path.expanduser(raw)))
        ok, motivo = _validar_docx(ruta)
        if ok:
            return str(ruta)

        print(f"  ERROR: {motivo}")
        if not ruta.exists() and ruta.parent.exists():
            docxs = sorted(ruta.parent.glob("*.docx"))
            if docxs:
                print(f"  Archivos .docx en {ruta.parent}:")
                for d in docxs[:10]:
                    print(f"    - {d.name}")


# ─────────────────────────────────────────────────────────────
# FASE 1: Parámetros de BD (mínimos para ejecutar queries)
# ─────────────────────────────────────────────────────────────

def solicitar_parametros_bd():
    """Solicita solo los parámetros necesarios para conectar y extraer datos."""

    print("\n" + "=" * 60)
    print("  A006 — GENERADOR DE INFORMES DE ANÁLISIS HIDRÁULICO")
    print("=" * 60)

    # 1. Verificar conexión
    print("\n[1/4] Verificando conexión a la base de datos...")
    if not test_connection():
        print("\nERROR: No se puede conectar a la base de datos.")
        print("Comprueba los parámetros en config.py y que el servidor esté accesible.")
        sys.exit(1)

    # 2. muni_id
    print("\n[2/4] Identificación del municipio")
    muni_id = input("  Introduce el muni_id: ").strip()
    nombre = get_municipio_nombre(muni_id)
    if not nombre:
        print(f"  ERROR: No se encontró ningún municipio con muni_id={muni_id}")
        sys.exit(1)
    print(f"  Municipio: {nombre}")

    # 3. Sectores
    print("\n[3/4] Sectores hidráulicos")
    num_sectores = int(input("  Número de sectores hidráulicos: ").strip())
    sector_ids = []
    for i in range(num_sectores):
        sid = input(f"    sector_id [{i + 1}/{num_sectores}]: ").strip()
        sector_ids.append(int(sid))
    print(f"  Sectores: {sector_ids}")

    # 4. result_id
    print("\n[4/4] Simulación EPANET")
    disponibles = get_result_ids_disponibles()
    if disponibles:
        print("  Simulaciones disponibles en la base de datos:")
        for r in disponibles:
            print(f"    - {r['result_id']}  ({r.get('num_timesteps', '?')} timesteps)")
    result_id = input("  Introduce el result_id a usar: ").strip()

    return {
        "muni_id": muni_id,
        "municipio_nombre": nombre,
        "num_sectores": num_sectores,
        "sector_ids": sector_ids,
        "result_id": result_id,
    }


# ─────────────────────────────────────────────────────────────
# FASE 2: Parámetros del documento (después de extraer datos)
# ─────────────────────────────────────────────────────────────

def solicitar_parametros_documento(params):
    """Solicita los parámetros del documento Word y campos manuales."""

    print("\n" + "=" * 60)
    print("  DATOS PARA EL DOCUMENTO")
    print("=" * 60)

    # 1. Plantilla Word
    print("\n[1/3] Plantilla Word")
    plantilla = pedir_ruta_plantilla()
    print(f"  Plantilla: {Path(plantilla).name}")
    params["plantilla"] = plantilla

    # 2. Campos manuales
    print("\n[2/3] Datos manuales del municipio")
    print("  (Introduce los valores o pulsa Enter para dejar vacío)")
    campos_manuales = {}
    for clave, descripcion in CAMPOS_MANUALES:
        valor = input(f"  {descripcion}: ").strip()
        campos_manuales[clave] = valor
    params["campos_manuales"] = campos_manuales

    # 3. Descripción de sectores
    print("\n[3/3] Descripción de la sectorización")
    print("  Escribe un párrafo descriptivo de los sectores hidráulicos del municipio.")
    print("  (Este texto se insertará en el marcador 'descripcion_sectores')")
    descripcion_sectores = input("  Texto: ").strip()
    params["descripcion_sectores"] = descripcion_sectores

    return params


# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────

def main():
    # Verificar que existe config.py
    if not Path("config.py").exists():
        print("ERROR: No se encuentra config.py")
        print("Copia config_example.py como config.py y rellena tus credenciales.")
        sys.exit(1)

    from config import OUTPUT_DIR
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

    # ── FASE 1: Parámetros de BD ──
    params = solicitar_parametros_bd()

    # ── FASE 2: Extracción de datos ──
    print("\n" + "=" * 60)
    print("  EXTRAYENDO DATOS DE LA BASE DE DATOS...")
    print("=" * 60)

    print("\n  Datos del municipio...")
    datos_muni = get_datos_municipio(params["muni_id"], params["sector_ids"])

    print("  Datos de sectores...")
    datos_sectores = get_datos_sectores(params["sector_ids"])

    print("  Resultados de simulación...")
    resultados = get_resultados_simulacion(params["result_id"], params["sector_ids"])

    print("\n  ✓ Datos extraídos correctamente.")
    print(f"    Longitud red: {datos_muni.get('longitud_red', '?')} km")
    print(f"    Nodos: {datos_muni.get('num_nodos', '?')}")
    print(f"    Arcos: {datos_muni.get('num_arcos', '?')}")
    print(f"    Depósitos: {datos_muni.get('num_depositos', '?')}")
    print(f"    Sectores con datos: {len(datos_sectores)}")

    # Detecta si existe ya un informe previo del municipio. Si existe, ofrece
    # saltarse las fases 3 y 4 y enviar el Word existente directamente a Claude.
    # Útil cuando ya generaste y editaste el Word a mano y solo quieres
    # (re)generar los textos narrativos.
    from word.properties import rellenar_docproperties
    from word.tables import rellenar_tablas
    from word.sectors import replicar_sectores
    from word.bookmarks import rellenar_marcadores
    import docx

    slug = re.sub(r'[<>:"/\\|?*]', '_', params['municipio_nombre'].upper())
    slug = re.sub(r'\s+', '_', slug).strip('_')
    slug = re.sub(r'_+', '_', slug)
    nombre_salida = f"A006_{slug}.docx"
    ruta_salida = Path(OUTPUT_DIR) / nombre_salida

    reutilizar_existente = False
    if ruta_salida.exists():
        print("\n" + "=" * 60)
        print(f"  Existe un informe previo: {ruta_salida.name}")
        print("=" * 60)
        print("  ¿Reutilizar el Word existente y saltar a la fase Claude?")
        print("    s = sí: conserva el Word actual (con tus ediciones manuales)")
        print("            y va directamente a generar los textos con Claude.")
        print("    n = no: regenera el Word desde plantilla (pierde tus ediciones)")
        print("            y pide plantilla + datos manuales.")
        respuesta = input("  Selecciona [s/n] (por defecto: s): ").strip().lower()
        reutilizar_existente = (respuesta != "n")

    if reutilizar_existente:
        ok, motivo = _validar_docx(ruta_salida)
        if not ok:
            print(f"\nERROR: el informe previo no se puede abrir:\n  {ruta_salida}")
            print(f"  Causa: {motivo}")
            sys.exit(1)
        print(f"\n  ✓ Reutilizando informe existente: {ruta_salida}")
        doc = None  # se abrirá en la fase Claude
    else:
        # ── FASE 3: Parámetros del documento ──
        params = solicitar_parametros_documento(params)

        # ── FASE 4: Generación del informe ──
        print("\n" + "=" * 60)
        print("  GENERANDO INFORME WORD...")
        print("=" * 60)

        plantilla_path = Path(params["plantilla"])
        ok, motivo = _validar_docx(plantilla_path)
        if not ok:
            print(f"\nERROR: no se puede abrir la plantilla:\n  {plantilla_path}")
            print(f"  Causa: {motivo}")
            sys.exit(1)
        doc = docx.Document(str(plantilla_path))

        print("\n  [1/4] Rellenando campos DOCPROPERTY...")
        rellenar_docproperties(doc, params, datos_muni, resultados)

        print("  [2/4] Rellenando tablas...")
        rellenar_tablas(doc, datos_muni, datos_sectores, resultados)

        print("  [3/4] Replicando bloques de sector...")
        replicar_sectores(doc, datos_sectores, resultados)

        # Insertar textos manuales del usuario
        from word.bookmarks import _find_bookmarks, _insert_text_at_bookmark
        bookmarks = _find_bookmarks(doc)
        for bm_name in ("descripcion_sectores",):
            texto = params.get(bm_name, "")
            if texto and bm_name in bookmarks:
                _insert_text_at_bookmark(bookmarks[bm_name], texto)

        print("  [4/4] Guardando documento...")
        doc.save(str(ruta_salida))

        print("\n" + "=" * 60)
        print(f"  ✓ Informe generado: {ruta_salida}")
        print("=" * 60)

    # ── FASE 5: Generación de textos con Claude (tras revisión manual) ──
    print("\n  Revisa y edita el documento en Word antes de continuar.")
    print(f"  Archivo: {ruta_salida}")
    generar_ai = input("\n  ¿Enviar a Claude para generar textos narrativos? (s/n): ").strip().lower()

    if generar_ai == "s":
        print("\n" + "=" * 60)
        print("  GENERANDO TEXTOS CON CLAUDE API...")
        print("=" * 60)

        # Reabrir el documento editado por el usuario
        doc = docx.Document(str(ruta_salida))

        print("\n  Generando textos narrativos...")
        rellenar_marcadores(doc, params, datos_muni, datos_sectores, resultados)

        doc.save(str(ruta_salida))

        print("\n" + "=" * 60)
        print(f"  ✓ Textos narrativos añadidos: {ruta_salida}")
        print("=" * 60 + "\n")
    else:
        print("\n  Ejecución finalizada. Los marcadores narrativos quedan vacíos.")
        print("  Puedes volver a ejecutar el script para generar los textos.\n")


if __name__ == "__main__":
    main()
